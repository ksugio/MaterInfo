from django.http import HttpResponse
from django.urls import reverse
from django.shortcuts import get_object_or_404
from django.template import Template, Context
from project.views import base, base_api, remote
from project.models import Project
from project.forms import ImportForm, CloneForm, TokenForm, SetRemoteForm, SearchForm
from ..models.reference import Reference
from ..models.article import Article, ArticleQueryset
from ..serializer import ReferenceSerializer
from .article import ArticleRemote
from Levenshtein import distance
from pypdf import PdfWriter
from docx import Document
from io import BytesIO
from zipfile import ZipFile
import datetime
import os
import io
import json
import bibtexparser

class AddView(base.AddView):
    model = Reference
    upper = Project
    fields = ('title', 'note', 'order', 'pagesize', 'template', 'startid')
    template_name ="project/default_add.html"

class ListView(base.ListView):
    model = Reference
    upper = Project
    template_name = "project/default_list.html"
    navigation = [['Add', 'reference:add'],
                  ['Import', 'reference:import'],
                  ['Clone', 'reference:clone'],
                  ['Search', 'reference:search']]

def ArticleSummary(upper):
    articles = Article.objects.filter(upper=upper)
    return {
        'narticles': len(articles)
    }

class DetailView(base.DetailView):
    model = Reference
    template_name = "reference/reference_detail.html"
    navigation = []

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        model = self.model.objects.get(pk=self.kwargs['pk'])
        url = reverse('reference:article_list',
                      kwargs={'pk': model.id, 'order': model.order, 'size':model.pagesize})
        nav_list = [base.Link('Article', url)]
        context['navigation_list'] = nav_list
        context['summary'] = ArticleSummary(model)
        if model.data:
            context['model_data'] = json.loads(model.data)
        return context

def RenderText(model):
    try:
        lines = []
        template = Template(model.template)
        for i, art in enumerate(ArticleQueryset(model, model.order)):
            context = Context({
                'id': i + model.startid,
                'type': art.get_type_display(),
                'key': art.key,
                'title': art.title,
                'author': art.author_display(),
                'journal': art.journal,
                'volume': art.volume,
                'number': art.number,
                'month': art.month,
                'year': art.year,
                'pages': art.pages,
                'cited': art.cited,
                'impact': art.impact,
                'booktitle': art.booktitle,
                'publisher': art.publisher,
                'address': art.address,
                'doi': art.doi,
                'url': art.url,
                'filename': art.file.name,
                'abstract': art.abstract,
                'note': art.note
            })
            line = template.render(context)
            line = line.replace('&amp;', '&')
            lines.append([line, art.id])
        return lines
    except:
        return []

def checkSimilar(lines):
    similar = []
    tot = len(lines)
    for i in range(tot):
        for j in range(i+1, tot):
            sim = 1 - distance(lines[i][0], lines[j][0]) / max(len(lines[i][0]), len(lines[j][0]))
            if sim > 0.9:
                similar.append([i, j, sim])
    return similar

def BibtexText(model):
    entries = []
    for art in ArticleQueryset(model, model.order):
        entries.append(art.bibentry())
    db = bibtexparser.bibdatabase.BibDatabase()
    db.entries = entries
    return bibtexparser.dumps(db)

class UpdateView(base.UpdateView):
    model = Reference
    fields = ('title', 'status', 'note', 'order', 'pagesize', 'template', 'startid')
    template_name = "project/default_update.html"

    def form_valid(self, form):
        model = form.save(commit=False)
        model.updated_by = self.request.user
        lines = RenderText(model)
        bibtex = BibtexText(model)
        params = {
            'lines': lines,
            'similar': checkSimilar(lines),
            'bibtex': bibtex
        }
        model.data = json.dumps(params)
        return super().form_valid(form)

class DeleteView(base.DeleteManagerView):
    model = Reference
    template_name = "project/default_delete.html"

class EditNoteView(base.MDEditView):
    model = Reference
    text_field = "note"
    template_name = "project/default_mdedit.html"

class MergePDFView(base.View):
    model = Reference

    def get(self, request, **kwargs):
        model = self.model.objects.get(pk=kwargs['pk'])
        merger = PdfWriter()
        for art in ArticleQueryset(model, model.order):
            if art.file is not None and os.path.splitext(art.file.name)[-1] == '.pdf':
                with art.file.open('rb') as f:
                    merger.append(f)
        buf = io.BytesIO()
        merger.write(buf)
        response = HttpResponse(buf.getvalue(), content_type='application/pdf')
        buf.close()
        merger.close()
        now = datetime.datetime.now()
        filename = 'MergePDF%s.pdf' % (now.strftime('%Y%m%d%H%M%S')[2:])
        response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'{}'.format(filename)
        return response

class BibtexView(base.View):
    model = Reference

    def test_func(self):
        model = get_object_or_404(self.model, unique=self.kwargs['unique'])
        return self.request.user in base.ProjectMember(model)

    def get(self, request, **kwargs):
        model = self.model.objects.get(unique=kwargs['unique'])
        data = json.loads(model.data)
        response = HttpResponse(data['bibtex'], content_type="text/plain")
        now = datetime.datetime.now()
        filename = 'Reference%s.bib' % (now.strftime('%Y%m%d%H%M%S')[2:])
        response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'{}'.format(filename)
        return response

class DocxView(base.View):
    model = Reference

    def get(self, request, **kwargs):
        model = self.model.objects.get(pk=kwargs['pk'])
        data = json.loads(model.data)
        document = Document()
        document.add_heading(model.title, 0)
        for line in data['lines']:
            document.add_paragraph(line[0])
        buf = io.BytesIO()
        document.save(buf)
        response = HttpResponse(buf.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        buf.close()
        now = datetime.datetime.now()
        filename = 'Reference%s.docx' % (now.strftime('%Y%m%d%H%M%S')[2:])
        response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'{}'.format(filename)
        return response

class DownloadZipView(base.View):
    model = Reference

    def get(self, request, **kwargs):
        model = self.model.objects.get(pk=kwargs['pk'])
        articles = Article.objects.filter(upper=model)
        buf = BytesIO()
        zipf = ZipFile(buf, 'w')
        data = []
        for art in articles:
            data.append(art.get_data())
            if art.file:
                filename = os.path.basename(art.file.name)
                zipf.writestr(filename, art.file.read())
        zipf.writestr('reference.json', json.dumps(data, indent=2))
        zipf.close()
        response = HttpResponse(buf.getvalue(), content_type='application/zip')
        buf.close()
        now = datetime.datetime.now()
        filename = 'Reference%s.zip' % (now.strftime('%Y%m%d%H%M%S')[2:])
        response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'{}'.format(filename)
        return response

class SearchView(base.SearchView):
    model = Reference
    upper = Project
    form_class = SearchForm
    template_name = "reference/reference_search.html"
    title = 'Reference Search'
    session_name = 'reference_search'
    back_name = 'reference:list'
    lower_items = [{'Search': 'reference.views.article.ArticleSearch'}]

# API
class AddAPIView(base_api.AddAPIView):
    upper = Project
    serializer_class = ReferenceSerializer

class ListAPIView(base_api.ListAPIView):
    model = Reference
    upper = Project
    serializer_class = ReferenceSerializer

class RetrieveAPIView(base_api.RetrieveAPIView):
    model = Reference
    serializer_class = ReferenceSerializer

class UpdateAPIView(base_api.UpdateAPIView):
    model = Reference
    serializer_class = ReferenceSerializer

class FileAPIView(base_api.FileAPIView):
    model = Reference
    attachment = False

class ReferenceRemote(remote.FileRemote):
    model = Reference
    add_name = 'reference:api_add'
    list_name = 'reference:api_list'
    retrieve_name = 'reference:api_retrieve'
    update_name = 'reference:api_update'
    file_fields_names = [('file', 'reference:api_file')]
    serializer_class = ReferenceSerializer
    lower_remote = [ArticleRemote]

class ImportView(remote.ImportView):
    model = Reference
    upper = Project
    form_class = ImportForm
    remote_class = ReferenceRemote
    template_name = "project/default_import.html"
    title = 'Reference Import'
    success_name = 'reference:list'
    view_name = 'reference:detail'
    hidden_lower = False
    default_lower = True

class CloneView(remote.CloneView):
    model = Reference
    form_class = CloneForm
    upper = Project
    remote_class = ReferenceRemote
    template_name = "project/default_clone.html"
    title = 'Reference Clone'
    success_name = 'reference:list'
    view_name = 'reference:detail'

class TokenView(remote.TokenView):
    model = Reference
    form_class = TokenForm
    success_names = ['reference:pull', 'reference:push']

class PullView(remote.PullView):
    model = Reference
    remote_class = ReferenceRemote
    success_name = 'reference:detail'
    fail_name = 'reference:token'

class PushView(remote.PushView):
    model = Reference
    remote_class = ReferenceRemote
    success_name = 'reference:detail'
    fail_name = 'reference:token'

class LogView(remote.LogView):
    model = Reference

class SetRemoteView(remote.SetRemoteView):
    model = Reference
    form_class = SetRemoteForm
    remote_class = ReferenceRemote
    title = 'Reference Set Remote'
    success_name = 'reference:detail'
    view_name = 'reference:detail'

class ClearRemoteView(remote.ClearRemoteView):
    model = Reference
    remote_class = ReferenceRemote
    success_name = 'reference:detail'
